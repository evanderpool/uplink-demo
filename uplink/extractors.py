"""File-type extractors.

Every extractor turns one source file into a list of Sections — a section
title plus its plain text. The chunker downstream never needs to know what
kind of file the text came from.

Supported today:
    .md / .markdown / .txt   stdlib
    .csv / .tsv              stdlib
    .pdf                     pypdf        (optional dependency)
    .docx                    python-docx  (optional dependency)
    .xlsx / .xlsm            openpyxl     (optional dependency)
    .xls                     xlrd         (optional dependency)

Optional dependencies degrade gracefully: if the library is missing the file
is skipped with a clear message instead of crashing the whole index run.
"""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from pathlib import Path

# Rows caps keep a runaway spreadsheet from flooding the index.
MAX_TABLE_ROWS = 2000
MAX_CELL_CHARS = 500


@dataclass
class Section:
    title: str
    text: str
    # For tabular data: a header line re-prepended to every chunk so rows
    # keep their column context after chunking.
    header: str = ""


@dataclass
class Extracted:
    sections: list[Section] = field(default_factory=list)
    title: str = ""
    warnings: list[str] = field(default_factory=list)


class ExtractorUnavailable(RuntimeError):
    """Raised when a file type needs a library that is not installed."""


SUPPORTED_EXTENSIONS = {
    ".md", ".markdown", ".txt",
    ".csv", ".tsv",
    ".pdf", ".docx",
    ".xlsx", ".xlsm", ".xls",
}

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def extract(path: Path) -> Extracted:
    """Dispatch on file extension. Raises ExtractorUnavailable or ValueError."""
    ext = path.suffix.lower()
    if ext in (".md", ".markdown"):
        return _extract_markdown(path)
    if ext == ".txt":
        return _extract_text(path)
    if ext in (".csv", ".tsv"):
        return _extract_delimited(path, "\t" if ext == ".tsv" else ",")
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(path)
    if ext == ".xls":
        return _extract_xls(path)
    raise ValueError(f"Unsupported file type: {ext}")


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_markdown(path: Path) -> Extracted:
    """Split a markdown file into sections at its headings."""
    out = Extracted(title=path.stem)
    current_title = path.stem
    buf: list[str] = []

    def flush() -> None:
        text = "\n".join(buf).strip()
        if text:
            out.sections.append(Section(title=current_title, text=text))
        buf.clear()

    in_fence = False
    for line in _read_text(path).splitlines():
        if line.lstrip().startswith(("```", "~~~")):
            in_fence = not in_fence
            buf.append(line)
            continue
        m = None if in_fence else _HEADING_RE.match(line)
        if m:
            flush()
            current_title = m.group(2).strip() or current_title
            if not out.title or out.title == path.stem:
                # First heading becomes the document title.
                if m.group(1) == "#":
                    out.title = current_title
        else:
            buf.append(line)
    flush()

    if not out.sections:
        out.warnings.append("empty document")
    return out


def _extract_text(path: Path) -> Extracted:
    text = _read_text(path).strip()
    out = Extracted(title=path.stem)
    if text:
        out.sections.append(Section(title=path.stem, text=text))
    else:
        out.warnings.append("empty document")
    return out


def _extract_delimited(path: Path, delimiter: str) -> Extracted:
    """CSV/TSV: header + rows rendered one line per row, ' | ' separated."""
    out = Extracted(title=path.stem)
    raw = _read_text(path)
    reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
    rows = []
    header = ""
    for i, row in enumerate(reader):
        cells = [c.strip()[:MAX_CELL_CHARS] for c in row]
        line = " | ".join(cells).strip()
        if not line.strip(" |"):
            continue
        if not header:
            header = line
            continue
        rows.append(line)
        if len(rows) >= MAX_TABLE_ROWS:
            out.warnings.append(f"truncated at {MAX_TABLE_ROWS} rows")
            break
    if header or rows:
        body = "\n".join([header] + rows) if header else "\n".join(rows)
        out.sections.append(Section(title=path.stem, text=body, header=header))
    else:
        out.warnings.append("empty document")
    return out


def _extract_pdf(path: Path) -> Extracted:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise ExtractorUnavailable("pypdf not installed (pip install pypdf)") from exc

    out = Extracted(title=path.stem)
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            out.warnings.append("encrypted PDF - skipped")
            return out
    meta_title = (reader.metadata.title or "").strip() if reader.metadata else ""
    if meta_title:
        out.title = meta_title
    for n, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            out.warnings.append(f"page {n}: extraction failed ({exc})")
            continue
        if text:
            out.sections.append(Section(title=f"Page {n}", text=text))
    if not out.sections:
        out.warnings.append("no extractable text (scanned/image PDF?)")
    return out


def _extract_docx(path: Path) -> Extracted:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise ExtractorUnavailable(
            "python-docx not installed (pip install python-docx)"
        ) from exc

    out = Extracted(title=path.stem)
    document = docx.Document(str(path))
    current_title = path.stem
    buf: list[str] = []

    def flush() -> None:
        text = "\n".join(buf).strip()
        if text:
            out.sections.append(Section(title=current_title, text=text))
        buf.clear()

    for para in document.paragraphs:
        style = (para.style.name or "") if para.style is not None else ""
        text = para.text.strip()
        if style.startswith("Heading") and text:
            flush()
            current_title = text
            if out.title == path.stem:
                out.title = text
        elif text:
            buf.append(text)
    flush()

    # Tables: one section per table, rows rendered like CSV lines.
    for t_i, table in enumerate(document.tables, start=1):
        lines = []
        for row in table.rows:
            cells = [c.text.strip()[:MAX_CELL_CHARS] for c in row.cells]
            line = " | ".join(cells).strip()
            if line.strip(" |"):
                lines.append(line)
            if len(lines) >= MAX_TABLE_ROWS:
                out.warnings.append(f"table {t_i} truncated at {MAX_TABLE_ROWS} rows")
                break
        if lines:
            header = lines[0] if len(lines) > 1 else ""
            out.sections.append(
                Section(title=f"Table {t_i}", text="\n".join(lines), header=header)
            )

    if not out.sections:
        out.warnings.append("empty document")
    return out


def _extract_xlsx(path: Path) -> Extracted:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise ExtractorUnavailable(
            "openpyxl not installed (pip install openpyxl)"
        ) from exc

    out = Extracted(title=path.stem)
    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        for sheet in wb.worksheets:
            lines: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v).strip()[:MAX_CELL_CHARS] for v in row]
                line = " | ".join(cells).strip()
                if line.strip(" |"):
                    lines.append(line)
                if len(lines) >= MAX_TABLE_ROWS:
                    out.warnings.append(
                        f"sheet '{sheet.title}' truncated at {MAX_TABLE_ROWS} rows"
                    )
                    break
            if lines:
                header = lines[0] if len(lines) > 1 else ""
                out.sections.append(
                    Section(title=f"Sheet: {sheet.title}", text="\n".join(lines), header=header)
                )
    finally:
        wb.close()

    if not out.sections:
        out.warnings.append("empty workbook")
    return out


def _extract_xls(path: Path) -> Extracted:
    """Legacy binary Excel (BIFF). Same shape as _extract_xlsx, via xlrd."""
    try:
        import xlrd
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise ExtractorUnavailable("xlrd not installed (pip install xlrd)") from exc

    out = Extracted(title=path.stem)
    try:
        # xlrd prints benign structural notes ("OLE2 inconsistency", odd file
        # size) to its logfile for real-world .xls exports — the file still
        # reads fine. Send them to a throwaway buffer so they don't flood the
        # console or a deploy log. Real problems still raise XLRDError below.
        wb = xlrd.open_workbook(str(path), logfile=io.StringIO())
    except xlrd.XLRDError as exc:
        # Encrypted or corrupt workbook: warn and move on, like encrypted PDFs.
        out.warnings.append(f"unreadable .xls - skipped ({exc})")
        return out
    try:
        for sheet in wb.sheets():
            lines: list[str] = []
            for r in range(sheet.nrows):
                cells: list[str] = []
                for c in range(sheet.ncols):
                    cell = sheet.cell(r, c)
                    value = cell.value
                    if cell.ctype == xlrd.XL_CELL_DATE:
                        try:
                            value = xlrd.xldate_as_datetime(
                                value, wb.datemode
                            ).isoformat(sep=" ")
                        except Exception:
                            pass
                    elif isinstance(value, float) and value.is_integer():
                        # xlrd reads every number as float; drop the '.0' noise.
                        value = int(value)
                    cells.append(str(value).strip()[:MAX_CELL_CHARS])
                line = " | ".join(cells).strip()
                if not line.strip(" |"):
                    continue
                lines.append(line)
                if len(lines) >= MAX_TABLE_ROWS:
                    out.warnings.append(
                        f"sheet '{sheet.name}' truncated at {MAX_TABLE_ROWS} rows"
                    )
                    break
            if lines:
                header = lines[0] if len(lines) > 1 else ""
                out.sections.append(
                    Section(title=f"Sheet: {sheet.name}", text="\n".join(lines), header=header)
                )
    finally:
        wb.release_resources()

    if not out.sections:
        out.warnings.append("empty workbook")
    return out
